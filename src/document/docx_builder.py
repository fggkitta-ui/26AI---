"""
Word 文档生成模块。

命名规则（赛题要求）：{用户的问题}_{时间}.docx
例如：最近3个月的上海区域内的充电桩招标信息都有哪些_202604071424.docx
"""

import re
from datetime import datetime
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt

from src.pipeline.summarizer import SummarizedItem

OUTPUT_DIR = Path("outputs")


def sanitize_filename(text: str) -> str:
    """去除 Windows/Unix 文件名中的非法字符。"""
    return re.sub(r'[\\/:*?"<>|]', "_", text)


def build_output_filename(user_query: str, now: datetime = None) -> Path:
    now = now or datetime.now()
    timestamp = now.strftime("%Y%m%d%H%M")
    filename = f"{sanitize_filename(user_query)}_{timestamp}.docx"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR / filename


def build_docx(user_query: str, items: List[SummarizedItem],
               time_range_desc: str, is_incremental: bool = False) -> Path:
    doc = Document()

    doc.add_heading(user_query, level=1)
    meta = doc.add_paragraph()
    meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
    doc.add_paragraph(f"检索条件：{time_range_desc}")
    if is_incremental:
        doc.add_paragraph("说明：本次为定时任务的增量推送，仅包含本次新增内容。")

    if not items:
        doc.add_paragraph("本次未检索到符合条件的新增招投标信息。")
    else:
        for idx, item in enumerate(items, start=1):
            doc.add_heading(f"{idx}. {item.title}", level=2)
            p = doc.add_paragraph()
            p.add_run(f"发布时间：{item.publish_time}").font.size = Pt(10)
            doc.add_paragraph(f"来源链接：{item.source_url}")
            doc.add_paragraph(f"核心内容：{item.core_content}")
            if item.attachment_urls:
                doc.add_paragraph("附件链接：" + "；".join(item.attachment_urls))

    output_path = build_output_filename(user_query)
    doc.save(output_path)
    return output_path
